"""
Output generator for Football Predictor
Creates various output formats (CSV, JSON, DOCX, PDF)
Optimized for Termux environment
"""

import pandas as pd
import json
import os
from datetime import datetime
from typing import List, Dict, Optional
import logging
import matplotlib.pyplot as plt
import matplotlib.patches as patches
from pathlib import Path

# Document generation
try:
    from docx import Document
    from docx.shared import Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logging.warning("python-docx not available")

try:
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib import colors
    from reportlab.lib.units import inch
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False
    logging.warning("reportlab not available")

from config_loader import ConfigLoader
from utils import setup_logging, format_probability, format_team_name, cleanup_temp_files

class OutputGenerator:
    """Generate various output formats for predictions"""
    
    def __init__(self):
        self.config = ConfigLoader()
        self.logger = setup_logging('output')
        
        # Configure matplotlib for Termux
        plt.switch_backend('Agg')  # Non-interactive backend
        plt.style.use('default')
        
        # Create output directory
        self.output_dir = Path('output')
        self.output_dir.mkdir(exist_ok=True)
    
    def save_to_csv(self, matches: List[Dict], filename: str = None) -> str:
        """Save predictions to CSV file"""
        if filename is None:
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            filename = f"predictions_{timestamp}.csv"
        
        filepath = self.output_dir / filename
        
        try:
            # Flatten match data for CSV
            csv_data = []
            
            for match in matches:
                row = {
                    'match_id': match.get('match_id', ''),
                    'date': match.get('match_date', '')[:16],  # Truncate microseconds
                    'league': match.get('league', ''),
                    'home_team': match.get('home_team', ''),
                    'away_team': match.get('away_team', ''),
                    'venue': match.get('venue', ''),
                    'status': match.get('status', 'scheduled')
                }
                
                # Predictions
                predictions = match.get('predictions', {})
                row.update({
                    'home_win_prob': predictions.get('home_win', 0),
                    'draw_prob': predictions.get('draw', 0),
                    'away_win_prob': predictions.get('away_win', 0),
                    'over_2_5_prob': predictions.get('over_2_5', 0)
                })
                
                # Meta information
                row.update({
                    'confidence': match.get('confidence', 0),
                    'factors_used': ', '.join(match.get('factors_used', [])),
                    'risks': ' | '.join(match.get('risks', [])),
                    'prediction_time': match.get('prediction_time', '')
                })
                
                # Weather data
                weather = match.get('weather', {})
                if weather:
                    row.update({
                        'temperature': weather.get('temperature', ''),
                        'conditions': weather.get('conditions', ''),
                        'wind_speed': weather.get('wind_speed', ''),
                        'precipitation': weather.get('precipitation', '')
                    })
                
                csv_data.append(row)
            
            # Create DataFrame and save
            df = pd.DataFrame(csv_data)
            df.to_csv(filepath, index=False)
            
            self.logger.info(f"Saved {len(matches)} predictions to {filepath}")
            return str(filepath)
            
        except Exception as e:
            self.logger.error(f"Error saving CSV: {e}")
            return ""
    
    def save_to_json(self, matches: List[Dict], filename: str = None) -> str:
        """Save predictions to JSON file"""
        if filename is None:
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            filename = f"predictions_{timestamp}.json"
        
        filepath = self.output_dir / filename
        
        try:
            # Prepare data with metadata
            output_data = {
                'metadata': {
                    'generated_at': datetime.now().isoformat(),
                    'total_matches': len(matches),
                    'version': '1.0'
                },
                'matches': matches
            }
            
            with open(filepath, 'w', encoding='utf-8') as f:
                json.dump(output_data, f, indent=2, default=str, ensure_ascii=False)
            
            self.logger.info(f"Saved {len(matches)} predictions to {filepath}")
            return str(filepath)
            
        except Exception as e:
            self.logger.error(f"Error saving JSON: {e}")
            return ""
    
    def create_prediction_chart(self, match: Dict, save_path: str = None) -> str:
        """Create prediction probability chart"""
        if save_path is None:
            match_id = match.get('match_id', 'unknown')
            save_path = f"temp_chart_{match_id}.png"
        
        try:
            predictions = match.get('predictions', {})
            
            # Data for chart
            outcomes = ['Home Win', 'Draw', 'Away Win']
            probabilities = [
                predictions.get('home_win', 0) * 100,
                predictions.get('draw', 0) * 100,
                predictions.get('away_win', 0) * 100
            ]
            
            colors_list = ['#4CAF50', '#FF9800', '#2196F3']  # Green, Orange, Blue
            
            # Create figure with small size for Termux
            fig, ax = plt.subplots(figsize=(8, 5))
            
            # Create horizontal bar chart
            bars = ax.barh(outcomes, probabilities, color=colors_list, alpha=0.8)
            
            # Customize chart
            ax.set_xlabel('Probability (%)', fontsize=10)
            ax.set_title(f"{format_team_name(match.get('home_team', ''))} vs {format_team_name(match.get('away_team', ''))}", 
                        fontsize=12, fontweight='bold')
            
            # Add percentage labels on bars
            for i, (bar, prob) in enumerate(zip(bars, probabilities)):
                width = bar.get_width()
                ax.text(width + 1, bar.get_y() + bar.get_height()/2, 
                       f'{prob:.1f}%', ha='left', va='center', fontweight='bold')
            
            # Set limits and grid
            ax.set_xlim(0, max(100, max(probabilities) * 1.2))
            ax.grid(axis='x', alpha=0.3)
            
            # Add confidence and over/under info
            confidence = match.get('confidence', 0) * 100
            over_25 = predictions.get('over_2_5', 0) * 100
            
            info_text = f"Confidence: {confidence:.1f}%\nOver 2.5 Goals: {over_25:.1f}%"
            ax.text(0.02, 0.98, info_text, transform=ax.transAxes, 
                   verticalalignment='top', fontsize=9, 
                   bbox=dict(boxstyle='round', facecolor='wheat', alpha=0.8))
            
            plt.tight_layout()
            plt.savefig(save_path, dpi=100, bbox_inches='tight')
            plt.close()
            
            return save_path
            
        except Exception as e:
            self.logger.error(f"Error creating chart: {e}")
            return ""
    
    def generate_word_report(self, matches: List[Dict], filename: str = None) -> str:
        """Generate Word document report"""
        if not DOCX_AVAILABLE:
            self.logger.warning("python-docx not available, skipping Word report")
            return ""
        
        if filename is None:
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            filename = f"football_predictions_{timestamp}.docx"
        
        filepath = self.output_dir / filename
        
        try:
            doc = Document()
            
            # Title
            title = doc.add_heading('Football Match Predictions', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Subtitle with date
            subtitle = doc.add_heading(f'Generated on {datetime.now().strftime("%B %d, %Y at %H:%M")}', level=2)
            subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            doc.add_page_break()
            
            # Summary section
            doc.add_heading('Summary', level=1)
            
            # Calculate summary stats
            total_matches = len(matches)
            high_confidence = sum(1 for m in matches if m.get('confidence', 0) > 0.7)
            leagues = set(m.get('league', '') for m in matches)
            
            summary_p = doc.add_paragraph()
            summary_p.add_run(f"Total Matches: ").bold = True
            summary_p.add_run(f"{total_matches}\n")
            summary_p.add_run(f"Leagues Covered: ").bold = True
            summary_p.add_run(f"{len(leagues)}\n")
            summary_p.add_run(f"High Confidence Predictions: ").bold = True
            summary_p.add_run(f"{high_confidence}\n")
            
            doc.add_page_break()
            
            # Individual match predictions
            doc.add_heading('Match Predictions', level=1)
            
            for i, match in enumerate(matches, 1):
                # Match header
                home_team = format_team_name(match.get('home_team', ''))
                away_team = format_team_name(match.get('away_team', ''))
                league = match.get('league', 'Unknown League')
                
                match_heading = doc.add_heading(f"{i}. {home_team} vs {away_team}", level=2)
                
                # Match details
                details_p = doc.add_paragraph()
                details_p.add_run("League: ").bold = True
                details_p.add_run(f"{league}\n")
                
                match_date = match.get('match_date', '')[:16] if match.get('match_date') else 'TBD'
                details_p.add_run("Date: ").bold = True
                details_p.add_run(f"{match_date}\n")
                
                venue = match.get('venue', 'Unknown')
                details_p.add_run("Venue: ").bold = True
                details_p.add_run(f"{venue}\n")
                
                # Predictions table
                predictions = match.get('predictions', {})
                
                pred_p = doc.add_paragraph()
                pred_p.add_run("Predictions:").bold = True
                
                pred_table = doc.add_table(rows=4, cols=2)
                pred_table.style = 'Table Grid'
                
                # Table data
                outcomes = [
                    ('Home Win', predictions.get('home_win', 0)),
                    ('Draw', predictions.get('draw', 0)),
                    ('Away Win', predictions.get('away_win', 0)),
                    ('Over 2.5 Goals', predictions.get('over_2_5', 0))
                ]
                
                for row_idx, (outcome, prob) in enumerate(outcomes):
                    row_cells = pred_table.rows[row_idx].cells
                    row_cells[0].text = outcome
                    row_cells[1].text = format_probability(prob)
                
                # Confidence and risks
                confidence = match.get('confidence', 0)
                conf_p = doc.add_paragraph()
                conf_p.add_run("Confidence: ").bold = True
                conf_p.add_run(format_probability(confidence))
                
                risks = match.get('risks', [])
                if risks:
                    risks_p = doc.add_paragraph()
                    risks_p.add_run("Risk Factors:").bold = True
                    for risk in risks[:3]:  # Limit to 3 risks
                        risks_p.add_run(f"\n• {risk}")
                
                # Add chart if possible
                chart_path = self.create_prediction_chart(match)
                if chart_path and os.path.exists(chart_path):
                    try:
                        doc.add_paragraph()  # Space
                        doc.add_picture(chart_path, width=Inches(6))
                        os.remove(chart_path)  # Clean up
                    except Exception:
                        pass  # Skip if chart fails
                
                # Page break except for last match
                if i < len(matches):
                    doc.add_page_break()
            
            # Footer
            doc.add_page_break()
            footer = doc.add_paragraph()
            footer.add_run("Generated by Football Predictor").italic = True
            footer.add_run(f"\n{datetime.now().isoformat()}")
            footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            doc.save(filepath)
            
            self.logger.info(f"Generated Word report: {filepath}")
            return str(filepath)
            
        except Exception as e:
            self.logger.error(f"Error generating Word report: {e}")
            return ""
    
    def generate_pdf_report(self, matches: List[Dict], filename: str = None) -> str:
        """Generate PDF report"""
        if not PDF_AVAILABLE:
            self.logger.warning("reportlab not available, skipping PDF report")
            return ""
        
        if filename is None:
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            filename = f"football_predictions_{timestamp}.pdf"
        
        filepath = self.output_dir / filename
        
        try:
            doc = SimpleDocTemplate(str(filepath), pagesize=A4)
            styles = getSampleStyleSheet()
            story = []
            
            # Title
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=24,
                spaceAfter=30,
                alignment=1  # Center
            )
            
            story.append(Paragraph("Football Match Predictions", title_style))
            story.append(Paragraph(f"Generated: {datetime.now().strftime('%B %d, %Y at %H:%M')}", 
                                 styles['Normal']))
            story.append(Spacer(1, 20))
            
            # Summary
            story.append(Paragraph("Summary", styles['Heading2']))
            
            total_matches = len(matches)
            high_confidence = sum(1 for m in matches if m.get('confidence', 0) > 0.7)
            leagues = set(m.get('league', '') for m in matches)
            
            summary_text = f"""
            <b>Total Matches:</b> {total_matches}<br/>
            <b>Leagues Covered:</b> {len(leagues)}<br/>
            <b>High Confidence Predictions:</b> {high_confidence}<br/>
            """
            
            story.append(Paragraph(summary_text, styles['Normal']))
            story.append(Spacer(1, 20))
            
            # Match predictions
            story.append(Paragraph("Match Predictions", styles['Heading2']))
            
            for i, match in enumerate(matches, 1):
                home_team = format_team_name(match.get('home_team', ''))
                away_team = format_team_name(match.get('away_team', ''))
                
                # Match header
                match_title = f"{i}. {home_team} vs {away_team}"
                story.append(Paragraph(match_title, styles['Heading3']))
                
                # Match details
                league = match.get('league', 'Unknown League')
                match_date = match.get('match_date', '')[:16] if match.get('match_date') else 'TBD'
                venue = match.get('venue', 'Unknown')
                
                details_text = f"""
                <b>League:</b> {league}<br/>
                <b>Date:</b> {match_date}<br/>
                <b>Venue:</b> {venue}<br/>
                """
                story.append(Paragraph(details_text, styles['Normal']))
                story.append(Spacer(1, 10))
                
                # Predictions table
                predictions = match.get('predictions', {})
                confidence = match.get('confidence', 0)
                
                table_data = [
                    ['Outcome', 'Probability'],
                    ['Home Win', format_probability(predictions.get('home_win', 0))],
                    ['Draw', format_probability(predictions.get('draw', 0))],
                    ['Away Win', format_probability(predictions.get('away_win', 0))],
                    ['Over 2.5 Goals', format_probability(predictions.get('over_2_5', 0))],
                    ['', ''],
                    ['Confidence', format_probability(confidence)]
                ]
                
                table = Table(table_data, colWidths=[2*inch, 1.5*inch])
                table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                    ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, 0), 12),
                    ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                    ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                    ('FONTNAME', (0, -1), (-1, -1), 'Helvetica-Bold'),
                    ('GRID', (0, 0), (-1, -1), 1, colors.black)
                ]))
                
                story.append(table)
                story.append(Spacer(1, 10))
                
                # Risks
                risks = match.get('risks', [])
                if risks:
                    risks_text = "<b>Risk Factors:</b><br/>" + "<br/>".join(f"• {risk}" for risk in risks[:3])
                    story.append(Paragraph(risks_text, styles['Normal']))
                
                story.append(Spacer(1, 20))
            
            # Build PDF
            doc.build(story)
            
            self.logger.info(f"Generated PDF report: {filepath}")
            return str(filepath)
            
        except Exception as e:
            self.logger.error(f"Error generating PDF report: {e}")
            return ""
    
    def create_summary_chart(self, matches: List[Dict], filename: str = None) -> str:
        """Create summary chart of all predictions"""
        if filename is None:
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            filename = f"summary_chart_{timestamp}.png"
        
        filepath = self.output_dir / filename
        
        try:
            # Aggregate data
            leagues = {}
            confidence_ranges = {'High (>70%)': 0, 'Medium (40-70%)': 0, 'Low (<40%)': 0}
            total_predictions = {'Home Win': 0, 'Draw': 0, 'Away Win': 0}
            
            for match in matches:
                # Count by league
                league = match.get('league', 'Unknown')[:20]  # Truncate long names
                leagues[league] = leagues.get(league, 0) + 1
                
                # Count by confidence
                confidence = match.get('confidence', 0)
                if confidence > 0.7:
                    confidence_ranges['High (>70%)'] += 1
                elif confidence > 0.4:
                    confidence_ranges['Medium (40-70%)'] += 1
                else:
                    confidence_ranges['Low (<40%)'] += 1
                
                # Count predictions
                predictions = match.get('predictions', {})
                max_prob = max(predictions.values()) if predictions else 0
                
                if predictions.get('home_win', 0) == max_prob:
                    total_predictions['Home Win'] += 1
                elif predictions.get('draw', 0) == max_prob:
                    total_predictions['Draw'] += 1
                else:
                    total_predictions['Away Win'] += 1
            
            # Create subplots
            fig, ((ax1, ax2), (ax3, ax4)) = plt.subplots(2, 2, figsize=(12, 10))
            fig.suptitle('Football Predictions Summary', fontsize=16, fontweight='bold')
            
            # 1. Leagues pie chart
            if leagues:
                ax1.pie(leagues.values(), labels=leagues.keys(), autopct='%1.1f%%', startangle=90)
                ax1.set_title('Matches by League')
            
            # 2. Confidence bar chart
            ax2.bar(confidence_ranges.keys(), confidence_ranges.values(), 
                   color=['green', 'orange', 'red'], alpha=0.7)
            ax2.set_title('Predictions by Confidence')
            ax2.set_ylabel('Number of Matches')
            plt.setp(ax2.xaxis.get_majorticklabels(), rotation=45)
            
            # 3. Prediction outcomes
            colors = ['#4CAF50', '#FF9800', '#2196F3']
            ax3.bar(total_predictions.keys(), total_predictions.values(), color=colors, alpha=0.7)
            ax3.set_title('Predicted Outcomes')
            ax3.set_ylabel('Number of Predictions')
            
            # 4. Stats text
            ax4.axis('off')
            stats_text = f"""
SUMMARY STATISTICS

Total Matches: {len(matches)}
Leagues: {len(leagues)}

High Confidence: {confidence_ranges['High (>70%)']}
Medium Confidence: {confidence_ranges['Medium (40-70%)']}
Low Confidence: {confidence_ranges['Low (<40%)']}

Home Wins Predicted: {total_predictions['Home Win']}
Draws Predicted: {total_predictions['Draw']}
Away Wins Predicted: {total_predictions['Away Win']}

Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}
            """
            ax4.text(0.1, 0.9, stats_text, transform=ax4.transAxes, fontsize=10,
                    verticalalignment='top', fontfamily='monospace',
                    bbox=dict(boxstyle='round', facecolor='lightgray', alpha=0.8))
            
            plt.tight_layout()
            plt.savefig(filepath, dpi=150, bbox_inches='tight')
            plt.close()
            
            self.logger.info(f"Created summary chart: {filepath}")
            return str(filepath)
            
        except Exception as e:
            self.logger.error(f"Error creating summary chart: {e}")
            return ""
    
    def generate_all_outputs(self, matches: List[Dict], formats: List[str] = None) -> Dict[str, str]:
        """Generate all requested output formats"""
        if formats is None:
            formats = self.config.get('output.formats', ['csv', 'json'])
        
        results = {}
        
        try:
            for fmt in formats:
                if fmt == 'csv':
                    results['csv'] = self.save_to_csv(matches)
                elif fmt == 'json':
                    results['json'] = self.save_to_json(matches)
                elif fmt == 'docx':
                    results['docx'] = self.generate_word_report(matches)
                elif fmt == 'pdf':
                    results['pdf'] = self.generate_pdf_report(matches)
                elif fmt == 'chart':
                    results['chart'] = self.create_summary_chart(matches)
            
            # Always create summary chart
            if 'chart' not in results:
                results['chart'] = self.create_summary_chart(matches)
            
            # Clean up temporary files
            cleanup_temp_files()
            
            self.logger.info(f"Generated outputs: {', '.join(results.keys())}")
            return results
            
        except Exception as e:
            self.logger.error(f"Error generating outputs: {e}")
            return results
    
    def create_league_summary(self, matches: List[Dict]) -> Dict[str, Dict]:
        """Create summary by league"""
        league_summary = {}
        
        for match in matches:
            league = match.get('league', 'Unknown')
            if league not in league_summary:
                league_summary[league] = {
                    'matches': 0,
                    'high_confidence': 0,
                    'avg_confidence': 0,
                    'predicted_outcomes': {'home': 0, 'draw': 0, 'away': 0}
                }
            
            summary = league_summary[league]
            summary['matches'] += 1
            
            confidence = match.get('confidence', 0)
            if confidence > 0.7:
                summary['high_confidence'] += 1
            
            # Running average of confidence
            summary['avg_confidence'] = (
                (summary['avg_confidence'] * (summary['matches'] - 1) + confidence) / 
                summary['matches']
            )
            
            # Count predicted outcomes
            predictions = match.get('predictions', {})
            if predictions:
                max_prob = max(predictions.values())
                if predictions.get('home_win', 0) == max_prob:
                    summary['predicted_outcomes']['home'] += 1
                elif predictions.get('draw', 0) == max_prob:
                    summary['predicted_outcomes']['draw'] += 1
                else:
                    summary['predicted_outcomes']['away'] += 1
        
        return league_summary
